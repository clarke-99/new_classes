#new test
from scipy.stats import skew, yeojohnson, shapiro, anderson, kstest, normaltest
import statsmodels.api as sm
from pathlib import Path
from statsmodels.stats.outliers_influence import variance_inflation_factor
import scipy.stats as stats
from sklearn.cluster import DBSCAN
from sklearn.metrics import silhouette_score, make_scorer
from sklearn.neighbors import NearestNeighbors
from sklearn.preprocessing import PolynomialFeatures, StandardScaler, RobustScaler
from sklearn.preprocessing import OrdinalEncoder, OneHotEncoder, LabelEncoder
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import os
import seaborn as sns
import scipy.stats
from pyod.models.mad import MAD
import warnings
from sklearn.model_selection import GridSearchCV, RandomizedSearchCV
from gower import gower_matrix
from sklearn.manifold import TSNE
from collections import OrderedDict
from docx import Document
from docx.shared import Inches

class Helper_Functions:
    def __init__(self, project):
        self.project = project
        self.figures = {}
    
    def convert_to_num(self, x, num_type = float):
        if float:
            while True:
                try:
                    x = num_type(x)
                    if isinstance(x, num_type):
                        break
                except ValueError:
                    x = input(f'{x} is invalid, please enter a float number.')
            return x

    def while_string_function(self, x):
        while True:
            if x.lower() in ('y', 'n'):
                break
            else:
                x = input('\nInvalid input - Y/N are only accepted inputs ')
        if x.lower() == 'y':
            return True
        else:
            return False

    def select_numeric_cols(self, data):
        
        data_to_drop = []

        for feature in data.columns:
            if data[feature].dtype == 'object':
                data_to_drop.append(feature)
            else:
                pass

        data = data.drop(columns = data_to_drop)
        return data

    def save_file(self, title, feature = None):
        if not isinstance(data,pd.DataFrame):
            data = pd.DataFrame(data)
        else:
            pass

        if feature:
            fig_of = f'{feature}'
        else:
            fig_of = 'full_dataset'

        home_dir = Path.home()
        file_path = os.path.join(home_dir, 'Desktop', 'Coding', 'projects', self.project, 'data_files', fig_of, f'{title}.csv')
        directory = os.path.dirname(file_path)
        os.makedirs(directory, exist_ok=True)
        data.to_csv(file_path, header=True)
        
    def reshape_array(self, arr, column_vector = True):
        if column_vector:
            print(f'\nBefore Reshaping: {arr.shape}')
            reshaped_arr = arr.values.reshape(-1, 1)
            print(f'After Reshaping: {reshaped_arr.shape}\n')
        else:
            print(f'\nBefore Reshaping: {arr.shape}')
            reshape_arr = arr.values.reshape(1, -1)
            print(f'After Reshaping: {reshaped_arr.shape}\n')
        return reshaped_arr
    
    def format_titles(self, title):
        print('\033[1m' + f'\n{title.title()}\n' + '\033[0m')

    def title_generator(self, title, feature = None, transformations = None, display_feat = False):
        if transformations:
            transformations = transformations.to_string().replace(',', ', ').title()
            title = f'{title.title()}' + f' After {transformations} transformations'.title()
        else:
            title = f'{title.title()} untransformed'.title()

        if feature and display_feat:
            title += f' for {feature.title()}'
        
        return title
    
    def check_assump_correct(self, assumptions, assumed_about = 'feature distribution', format = 'feature:distribution'):
        self.format_titles(f'check {assumed_about}')
        print(assumptions)
        assump_check = input(f'Do the above assumptions about {assumed_about} seem correct? Y/N ')
        new_list = []

        while not self.while_string_function(assump_check):
            features = list(assumptions.keys())
            print(f'\nPlease type in the incorrect feature out of {features} and the correct distribution assumption out of likley, potentially and unlikely')
            new_type = input(f'You should input your selection in the following format {format}, {format}... ')
            print('\n')
            num_changed = 0

            if ',' in new_type:
                new_list = new_type.strip().split(',')
            else: 
                new_list = [new_type.strip()]

            for elements in new_list:
                if ':' in elements:
                    feature, dist_assumption = elements.strip().split(':')
                    dist_assumption = dist_assumption.title() + ' Normal'
            
                    if feature in assumptions:
                        assumptions[feature] = dist_assumption
                        num_changed += 1
                        print(f'{feature} is {assumptions[feature]}')
                    else:
                        print(f'\n{feature} not found in dataset')
                else:
                    print(f'\nInvalid input format expected \'feature:distribution\', got {elements}')

            if num_changed == 0:
                print('No values changed, please try again.')  
            elif num_changed == len(new_list):
                print('\n')
                print(assumptions)
                assump_check = input('Are these correct? Y/N ')
            else:
                print(f'\nNumber of features ({num_changed}) changed does not match length of input values ({len(new_list)})')
            
        return assumptions
            
    def generate_report(self, feature = None, transformations=None):

        stat_summary = (round(self.data.describe(percentiles=[.25, .5, .75]).T, 2))


        doc = Document()

        if not transformations:
            title = 'For Untransformed Data'
        else:
            title = transformations.to_string().replace(' ', ', ').title()
            title = f'After {title}'


        doc.add_heading(f'Statistical Analysis Report {title}', level=1)

        doc.add_heading('Summary Statistics', level=2)
        doc.add_paragraph(stat_summary.to_string())
        for feature in self.data.columns:
            mode_string = self.stats.get_modes(feature)
            doc.add_paragraph(mode_string)

        # Add normality check results
        doc.add_heading('Normality Test Results', level=2)
        norm_results, skew_kurtosis = self.stas.statistical_analysis(feature, transformations)
        for test, result in norm_results.items():
            doc.add_paragraph(f"{test.replace('_', ' ').title()}: {'Normal' if result else 'Not Normal'}")

        # Add figures
        doc.add_heading('Figures', level=2)

        # for fig in self.figures:
        #     doc.add_heading(fig["title"], level=3)
        #     doc.add_picture(fig["file_path"], width=Inches(5.5))

        # Save the document
        home_dir = Path.home()
        file_path = os.path.join(home_dir, 'Desktop', 'Coding', 'projects', self.project, 'reports', 'statistical_analysis_report.docx')
        directory = os.path.dirname(file_path)
        os.makedirs(directory, exist_ok=True)
        doc.save(file_path)

        print(f'Report saved at {file_path}')

class Figure_Manager:
    def __init__(self, project):
        self.project = project
        self.figures = {}
        self.helper_functions = Helper_Functions(project)

    def make_heatmap(self, data, type= 'dataset_correlation', feature1 = None, feature2=None, transformations = None, data_types = None):

        encode = Pre_Processing(data)
        encoded_data, encoded_dict = encode.encoder(data_types)     
        
        if not feature1 and not feature2:
            if 'ordinal' in data_types.values():
                print('\nOrdinal Variables Present')
                data = encoded_data  
            else:
                print('\nNo Ordinal Variables - Selecting only numerical data')
                data = data.select_dtypes(include=(float, int))

        if not isinstance(data, pd.DataFrame):
            print(f"Current data type: {type(data)}")
            raise ValueError("Data should be a pandas DataFrame")

        if type == 'dataset_correlation':
            xlabel = 'Features'
            ylabel = 'Features'
            title = 'heatmap of feature correlation matrix'
            data = data.corr()
            fig_of = 'dataset_corr'

        elif type == 'dataset_missing':
            xlabel= 'Features'
            ylabel= 'Index'
            title= 'heatmap of missing values'
            fig_of = 'dataset_missing'

        # elif type == 'feat_correlation':
        #     xlabel= feature2
        #     ylabel= feature1
        #     title= f'correlation of values between {feature1} and {feature2}'
        #     data = data.corr()
        #     if feature1 is not None and feature2 is not None:
        #         corr_value = data[[feature1, feature2]].corr().iloc[0, 1]
        #         heatmap_data = pd.DataFrame([[corr_value]], columns=[feature2], index=[feature1])
        #         xlabel = feature2
        #         ylabel = feature1
        #         title = f'correlation of values between {feature1} and {feature2}'
        #         fig_of = f'corr_{feature1}_{feature2}'
        #     else:
        #         raise ValueError("Both feature1 and feature2 must be provided for feat_correlation type")

        else:
            if feature1 is not None:
                missing_data = data[[feature1]].isnull().astype(int)
                xlabel = 'Index'
                ylabel = feature1
                title = f'heatmap of missing values in {feature1}'
                fig_of = f'{feature1}_missing'
                heatmap_data = missing_data
            else:
                raise ValueError("feature1 must be provided for feature_missing type")      

        num_features = data.shape[1]
        figsize = (min(10 + num_features * 0.2, 20), min(8 + num_features * 0.15, 15))
        fontsize = min(12, max(8, 120 // num_features))
        plt.figure(figsize=figsize)
        sns.heatmap(data, annot=True, cmap='coolwarm', fmt='.2f', annot_kws={"size": fontsize}, cbar=False)
        plt.xlabel(xlabel)
        plt.ylabel(ylabel)
        title = self.helper_functions.title_generator(title, transformations)
        plt.title(title)
        self.save_fig(fig_type='heatmaps', fig_of = fig_of, feature = feature1, transformations=transformations)
        
        return encoded_data   
    
    def make_q_q(self, data, feature, transformations = None):
        fig_type = 'qq_plots'
        fig_of = feature
        plt.figure(figsize=(12, 12))
        stats.probplot(data, dist="norm", plot=plt)
        plt.xlabel('Theoretical Quantiles for Normally Distributed Data')
        plt.ylabel('Actual Quantiles')
        title = self.helper_functions.title_generator(f'Theoretical vs Actual Quantiles for {feature}', feature, transformations)
        plt.title(title)
        return self.save_fig(fig_type=fig_type, fig_of=fig_of, feature=feature, transformations=transformations)

    def make_bar(self, data, feature = None, transformations = None):
        plt.figure(figsize=(12, 12))
        plt.xlabel(feature)
        plt.ylabel('Count')
        data = pd.DataFrame(data=data, columns=[feature])
        sns.countplot(x=feature, data=data)
        title = self.helper_functions.title_generator(f'bar chart of {feature}', feature, transformations)
        plt.title(title)
        self.save_fig(fig_type='bar_charts', fig_of=title, feature=feature, transformations=transformations)

    def make_histogram(self, data, feature, transformations = None):
        title = self.helper_functions.title_generator('Histogram', transformations)
        
        iqr = np.percentile(data, 75) - np.percentile(data, 25)
        bin_width_fd = 2 * iqr / len(data)**(1/3)
        num_bins = int((np.max(data) - np.min(data)) / bin_width_fd) #using Freedman-Diaconis Rule for num bins

        plt.figure(figsize=(12,12))
        n, bins, _ = plt.hist(data, bins=num_bins, density=True, edgecolor='black', linewidth=1.5, alpha=0.7, label=feature)
        
        title = self.helper_functions.title_generator(f'histogram of {feature}', feature, transformations)
        title += ' with normal distribution overlaid'.title()
        fig_type = 'histograms'
        fig_of = f'hist_{feature}'

        # Get mean and standard deviation of the data
        mu, sigma = np.mean(data), np.std(data)

        # Plot normal distribution curve
        xmin, xmax = plt.xlim()
        x = np.linspace(xmin, xmax, 100)
        p = stats.norm.pdf(x, mu, sigma)
        scaling_factor = np.sum(n * np.diff(bins))

        plt.plot(x, p * scaling_factor, linewidth=2, color = 'orange', label='Normal Distribution')
        plt.title(title)
        plt.xlabel(feature.title())
        plt.ylabel('Frequency')
        plt.legend()
        self.save_fig(fig_type=fig_type, fig_of=fig_of, feature=feature, transformations=transformations)

    def make_box_plot(self, data, feature = None, cat_var = None):
        ylabel = []
        data_to_drop = []

        if cat_var:
            title = f'boxplot of features for {cat_var}'.title()
            fig_of = f'{cat_var}'
        else:
            title = 'boxplot of features'.title()
            fig_of = 'dataset'

        for feature in data.columns:
            if data[feature].dtype == 'object':
                data_to_drop.append(feature)
            else:
                ylabel.append(feature)
        yticks = [i+1 for i in range(len(ylabel))]

        data = data.drop(columns=data_to_drop)
        plt.figure(figsize=(12,12))
        plt.boxplot(data, vert=False, patch_artist=True)
        plt.yticks(yticks, ylabel)
        plt.xlabel('Values')
        plt.title(title)
        self.save_fig(fig_type='box_plots', fig_of=fig_of, cat_feature=cat_var, feature=feature)


    def make_violin_plots(self, data, transformations=None):
        cat_features = data.select_dtypes(include=(object))
        num_features = data.select_dtypes(include=(float, int))

        for feature in num_features:
            plt.figure(figsize=(12,12))
            sns.violinplot(y=data[feature])
            title = self.helper_functions.title_generator(f'violin plot of {feature}', feature=feature, transformations=transformations)
            plt.title(title)
            self.save_fig('mono_var_violins', fig_of=feature, transformations=transformations)

        if not cat_features.empty:
            for cat_feature in cat_features:
                for feature in num_features:
                    plt.figure(figsize=(12,12))
                    sns.violinplot(x=data[cat_feature], y=data[feature], data = data[[feature]])
                    title = self.helper_functions.title_generator(f'violin plot of {feature} grouped by {cat_feature}', transformations=transformations)
                    plt.title(title)
                    self.save_fig('bi_var_violins', fig_of=f'{feature}_{cat_feature}', transformations=transformations)

        if len(cat_features.columns) > 1: 
            for i, cat_feature1 in enumerate(cat_features):
                for cat_feature2 in cat_features[i+1:]:
                    for feature in num_features:
                        plt.figure(figsize=(12, 12))
                        sns.violinplot(x=data[cat_feature1], y=data[feature], hue=data[cat_feature2], split=True)
                        plt.title(self.helper_functions.title_generator(f'Violin plot of {feature} grouped by {cat_feature1} and {cat_feature2}', transformations=transformations))
                        self.save_fig('multi_var_violins', fig_of=f'{feature}_{cat_feature1}_{cat_feature2}', transformations=transformations)

    def save_fig(self, fig_type, fig_of, feature = None, cat_feature = None, transformations = None):
        if transformations:
            transforms = transformations.to_string().replace(', ', '_').lower()
            fig_of += f'_{transforms}'
        else:
            fig_of += '_untransformed'

        fig_of = fig_of.replace('\'', '').replace('[', '').replace(']', '').replace(', ', '_').lower().replace(' ', '_')
        
        fig_of=fig_of.lower().replace(' ', '_')
        home_dir = Path.home()
        fig_type = fig_type.lower().replace(' ', '_')
        #file_path = os.path.join(home_dir, 'Desktop', 'Coding', 'projects', self.project, 'figures', fig_of, fig_type)
        file_path = os.path.join(home_dir, 'Desktop', 'Coding', 'projects', self.project, 'figures', fig_type, fig_of)

        directory = os.path.dirname(file_path)
        os.makedirs(directory, exist_ok=True)
        #self.figures[file_name] = plt.gcf()
        plt.savefig(file_path, bbox_inches='tight') 
        plt.close()

    def create_figures(self, data, feature = None, transformations = None):

        if feature and transformations: 
            title = f'For {feature.title()} after {transformations}'
            self.make_q_q(data, title, feature)
            self.make_histogram(data, feature, transformations)
        elif feature and not transformations:
            title = f'For {feature.title()}'
            self.make_q_q(data, title, feature)
        elif transformations and not feature:
            title = f'For Dataset After {transformations}'
            self.make_q_q(data, title)
        else:
            title = 'For Dataset'
            self.make_q_q(data, title)

class Statistics:
    def __init__(self, data, project):
        self.data = data
        self.figure_manager = Figure_Manager(project)
        self.helper_functions = Helper_Functions(project)
        self.original_norm_res = {}
    
    def determine_data_type(self, data):
        self.helper_functions.format_titles('determining data types')

        print('Datetime - Data that refers to time periods.')
        print('Interval - Continuous data type without a true 0, e.g temp, can be negative.')
        print('Ratio - Continuous data type with a true 0, e.g. height, weight, cannot be negative.')
        print('Ordinal - Categorical data type, refers to a category, e.g. how satisfied a customer is, can be represented numerically (ordinal encoding).')
        print('Nominal - Categorical data type without clear ordering, e.g. species (binary or one hot encoding).\n')
        data_res = {}
        data_types = ['datetime', 'interval', 'ratio', 'ordinal', 'nominal']
        
        for feature in data.columns:
            
            if data[feature].dtype in ['int64', 'float64']:
                if (data[feature] >= 0).all():
                    # Check if it's a date-like column
                    if pd.api.types.is_datetime64_any_dtype(data[feature]):
                        data_type = 'datetime'
                    else:
                        data_type = 'ratio'
                else:
                    data_type = 'ordinal' if abs(data[feature]).min() < 2 else 'interval'  # heuristic filtering for data type

            else:
                data_type = 'nominal'
                
            print(f"The data type for '{feature}' is '{data_type}'")
            data_res[feature] = data_type

        check_data = input('Are the assumed data types correct? Y/N ')
        while not self.helper_functions.while_string_function(check_data):
            selected_feature = False
            while not selected_feature:
                select_feature = input('Which feature is incorrect? ')
                if select_feature in data_res.keys():
                    values = ', '.join(data_types)
                    selected_feature = True
                else:
                    keys = str([key for key in data_res])[1:-1]
                    keys = keys.strip('\'')
                    select_feature = print(f'\n{select_feature} is an invalid selection. Please choose from {keys}\n')

            new_data_type = False
            while not new_data_type:
                new_data_type = input(f'What data type should it be out of {values}? ')
                if new_data_type in values:
                    data_res[select_feature] = new_data_type
                    new_data_type = True
                else:
                    print(f'{new_data_type} is an invalid input please select from {values}')
                    new_data_type = False

            check_data = input(f'Are {data_res} correct? ')
            self.helper_functions.while_string_function(check_data)

        return data_res

    def get_modes(self, feature):
        mode_values = self.data[feature].mode()
        mode_count = self.data[feature].eq(mode_values[0]).sum()
        total = self.data[feature].count()
        mode_percent = round((mode_count/total) * 100, 2)

        if len(mode_values.values) > 1:
            mode_modes = 'modes'
        else:
            mode_modes = 'mode'
                
        if len(mode_values.values) == self.data.shape[1]:
            mode_string = '\n{feature} has no modes - all values are unique\n'
        else:
            mode_string = '\033[1m' + f'{feature} {mode_modes}:' + '\033[0m' + f' {mode_values.values} = {mode_percent}% of values'
        return mode_string

    def normality_checks(self, data, skew_kurtosis, transformations = None): 

        title = 'checking normality of data'
        if transformations:
            title += f'after {transformations}'
            
        results = {}
        data = self.helper_functions.select_numeric_cols(self.data)
        n = len(data)
        
        if n <= 50:
            size = 'Small'
        elif n>50 and n <= 300:
            size = 'Medium'
        else:
            size = 'Large'

        self.helper_functions.format_titles(title)
        print(f'Dataset is {size}')

        def shapiro_wilks(data):
            stat, p_value = shapiro(data)
            return p_value > 0.05
        
        def andersond_darling(data):
            result = anderson(data)
            return result.statistic < result.critical_values[2]
        
        def kolmogorov_smirnov(data):
            stat, p_value = kstest(data, 'norm')
            return p_value > 0.05
        
        def lilliefors(data):
            stat, p_value = sm.stats.lilliefors(data)
            return p_value > 0.05
        
        def dagostino(data):
            stat, p_value = normaltest(data)
            return p_value > 0.05
        
        def z_skew(z_skew_score, size):
            if 'small' in size.lower():
                return z_skew_score <= 1.96
            elif 'medium' in size.lower():
                return z_skew_score <= 3.29
            else:
                return z_skew_score

        def z_kurtosis(z_kurtosis_score, size):
            if 'small' in size.lower():
                return z_kurtosis_score <= 1.96
            elif 'medium' in size.lower():
                return z_kurtosis_score <= 3.29
            else:
                return z_skew_score
        
        for feature in data.columns:
            results[feature] = {}
            skew = skew_kurtosis[feature]['skew']
            kurtosis = skew_kurtosis[feature]['kurtosis']
            kurtosis_excess = float(kurtosis) - 3
            skew_se = np.sqrt(6/n)
            kurtosis_se = np.sqrt(24/n)

            z_skew_score = abs(skew/skew_se)
            z_kurtosis_score = abs(kurtosis_excess/kurtosis_se)

            results[feature]['skew'] = skew
            results[feature]['z_skew'] = z_skew(z_skew_score=z_skew_score, size=size)
            results[feature]['kurtosis_excess'] = kurtosis_excess
            results[feature]['z_kurtosis'] = z_kurtosis(z_kurtosis_score=z_kurtosis_score, size=size)

            if n <= 50:
            
                results[feature][shapiro_wilks.__name__] = shapiro_wilks(data[feature])
                results[feature][andersond_darling.__name__] = andersond_darling(data[feature])

            elif n>50 and n<300:
            
                results[feature][shapiro_wilks.__name__] = shapiro_wilks(data[feature])
                results[feature][andersond_darling.__name__] = andersond_darling(data[feature])
                results[feature][lilliefors.__name__] = lilliefors(data[feature])
                results[feature][dagostino.__name__] = dagostino(data[feature])
                
            else:
                results[feature][shapiro_wilks.__name__] = shapiro_wilks(data[feature])
                results[feature][andersond_darling.__name__] = andersond_darling(data[feature])
                results[feature][lilliefors.__name__] = lilliefors(data[feature])
                results[feature][dagostino.__name__] = dagostino(data[feature])
                results[feature][kolmogorov_smirnov.__name__] = kolmogorov_smirnov(data[feature])
                
        return results

    def skew_kurtosis(self, data):
        skew_kurtosis = {}
        for feature in data.columns:
            skew_kurtosis[feature] = {}
            skew_kurtosis[feature]['skew'] = round(data[feature].skew(), 3)
            skew_kurtosis[feature]['kurtosis'] = round(data[feature].kurtosis(), 3)
        return skew_kurtosis
        
    def determine_normality(self, data, results, initial_analysis = True, transformations = None):
        n = len(data)
        final_result = {}

        if transformations:
            feature += f'after {transformations}'

        for feature in data.columns:
            self.helper_functions.format_titles(f'{feature} final results:')
            
            for test, result in results[feature].items():
                
                if result == True:
                    print(f'{feature} normally distributed according to {test} tests')

            if n < 300:
                if results[feature]['z_skew'] and results[feature]['z_kurtosis']:
                    print(f'{feature} distribution may be normal based on skewness and kurtosis')
                    final_result[feature] = 'likely normal'.title()

                elif results[feature]['z_skew'] or results[feature]['z_kurtosis']:
                    if results[feature]['z_skew']:
                        test_type = 'skewness'
                        excess = 'kurtosis'
                    else:
                        test_type = 'kurtosis'
                        excess = 'skewness'
                    print(f'{feature} distribution may be normal based on {test_type} but has excess {excess}')
                    final_result[feature] = 'potentially normal'.title()
                else:
                    print(f'{feature} distribution unlikely to be normal based on skewness or kurtosis')
                    final_result[feature] = 'unlikely normal'.title()


            else:
                if abs(results[feature]['skew']) <= 2 and abs(results[feature]['kurtosis_excess']) <= 4:
                    print(f'{feature} distribution likely to be normal based on skewness and kurtosis')
                    final_result[feature] = 'likely normal'.title()
                elif abs(results[feature]['skew']) <= 2 or abs(results[feature]['kurtosis_excess']) <= 4:
                    if results[feature]['z_skew']:
                        test_type = 'skewness'
                    else:
                        test_type = 'kurtosis'
                    print(f'{feature} distribution may be normal based on {test_type}')
                    final_result[feature] = 'potentially normal'.title()
                else:
                    print(f'{feature} distribution unlikely to be normal based on skewness or kurtosis')
                    final_result[feature] = 'unlikely normal'.title()


        self.helper_functions.format_titles('summary of results')
        for feat, res in final_result.items():
            print(f'{feat} is {res}')
        
        if not initial_analysis:
            feats_to_transform = []
            check_normal = input('\nBased on test results and visual inspection of figures, are these distributions normal? Y/N ')
            while not self.helper_functions.while_string_function(check_normal):
                print('\nIf you select all, all features will be transformed')
                transform_data = input('Which feature is not normal? ')
                if transform_data.lower() == 'all':
                    feats_to_transform = [key for key in final_result]
                    break
                elif transform_data in final_result.keys():
                    feats_to_transform.append(transform_data)
                    check_normal = input('Are all other\'s correct? Y/N ')
                    print('\n')
                else:
                    keys = str([key for key in final_result])[1:-1]
                    keys = keys.strip('\'')
                    transform_data = print(f'\n{transform_data} is an invalid selection. Please choose from {keys}\n')

            print(f'\nFeatures ready to transform: {str(feats_to_transform)}')
            return feats_to_transform, final_result
        
        else:
            return final_result
        
    def assess_transforms(self, results_dict, transformations = None):
        if not transformations:
            self.original_norm_res = results_dict
        else:
            for feature, results in results_dict.items():
                for norm_test, norm_result in results_dict.items():
                    pass
        # need to store results each time, and then assess how transforms have affected the distributions
        pass


    # def generate_report(self, feature = None, transformations=None):

        # stat_summary = (round(self.data.describe(percentiles=[.25, .5, .75]).T, 2))


        # doc = Document()

        # if not transformations:
        #     title = 'For Untransformed Data'
        # else:
        #     title = title = transformations.to_string().replace(' ', ', ').title()
        #     title = f'After {title}'


        # doc.add_heading(f'Statistical Analysis Report {title}', level=1)

        # doc.add_heading('Summary Statistics', level=2)
        # doc.add_paragraph(stat_summary.to_string())
        # for feature in self.data.columns:
        #     mode_string = self.get_modes(feature)
        #     doc.add_paragraph(mode_string)

        # # Add normality check results
        # doc.add_heading('Normality Test Results', level=2)
        # norm_results, skew_kurtosis = self.statistical_analysis(feature, transformations)
        # for test, result in norm_results.items():
        #     doc.add_paragraph(f"{test.replace('_', ' ').title()}: {'Normal' if result else 'Not Normal'}")

        # # Add figures
        # doc.add_heading('Figures', level=2)
        # for fig in self.figures:
        #     doc.add_heading(fig["title"], level=3)
        #     doc.add_picture(fig["file_path"], width=Inches(5.5))

        # # Save the document
        # home_dir = Path.home()
        # file_path = os.path.join(home_dir, 'Desktop', 'Coding', 'projects', self.project, 'reports', 'statistical_analysis_report.docx')
        # directory = os.path.dirname(file_path)
        # os.makedirs(directory, exist_ok=True)
        # doc.save(file_path)

        # print(f'Report saved at {file_path}')

class Pre_Processing:
    def __init__(self, data, project = None, target = None):
        #defining important variables 
        self.data = data.copy()
        self.project = project
        self.target = target 
        self.helper_functions = Helper_Functions(self.project)
        self.stats = Statistics(data=data, project=project)
        
        #defining encoders so that the labels can be decoded outside the encoder function
        self.label_encoder = LabelEncoder()
        self.one_hot_encoder = OneHotEncoder()
        self.categories = None
        self.ordinal_encoder = None
        
        #dictionary of dictionaries to store information about features, encoder type, and data
        self.encoded_feats = {}
    
    def ordinal_mapping(self, feature, ordinal_mapping):
        #create custom order
        sorted_ordinal_mapping = OrderedDict(sorted(ordinal_mapping.items())) #sort the dictionary of values and objects based on value
    
        print('\nMapping values:')

        print(sorted_ordinal_mapping)
        
        ordinal_mapping = list(sorted_ordinal_mapping.values()) #create list of values
        print('\nMapping list:')
        print(ordinal_mapping)
        categories = [ordinal_mapping] #create list of list to be passed to encoder
        print(categories)
        self.ordinal_encoder = OrdinalEncoder(categories=categories) #initialise encoder with the ordinal mapping as a list of lists
        
        if not self.ordinal_encoder:
            print('Encoder not intialised')
        else:
            pass

        #restructure data and fit encoder
        data_to_encode = self.helper_functions.reshape_array(self.data[feature]) 
        encoded_data = self.ordinal_encoder.fit_transform(data_to_encode) 
        #print(encoded_data)
        self.encoded_feats[self.ordinal_encoder] = {feature: encoded_data}
        encoder_mapping = self.ordinal_encoder.categories_[0]
        if encoder_mapping.tolist() == ordinal_mapping: #check encoding was successful
            print('Custom mapping successful')
        else:
            print('Mapping failed')
            print(f'Desired Map: {ordinal_mapping}')
    
        print(f'Custom mapping: {self.ordinal_encoder.categories_[0]}')
        return encoded_data
            
    def decoder(self, feature):
        decoded_data = None
        for encoder, encoded_dict in self.encoded_feats.items():
            if feature in encoded_dict:
                encoded_data = encoded_dict[feature]
                if encoder == self.label_encoder:
                    decoded_data = self.label_encoder.inverse_transform(encoded_data.reshape(-1, 1))
                elif encoder == self.one_hot_encoder:
                    decoded_data = self.one_hot_encoder.inverse_transform(encoded_data)
                elif encoder == self.ordinal_encoder:
                    decoded_data = self.ordinal_encoder.inverse_transform(encoded_data)
                break
        return decoded_data
        
    def encoder(self, data_types):
        data = self.data.copy()
        for feature in self.data.columns:
            if data_types[feature] == 'nominal' or data_types[feature] == 'ordinal':
                title = 'Encoded Dataframe:'
                cat_categories = set(data[feature])
                self.categories = list(cat_categories)
                if len(self.categories) == 1:
                    print('Single value cannot be encoded.')
                elif len(self.categories) == 2: 
                    encoded_data = self.label_encoder.fit_transform(data[feature])
                    data[feature] = encoded_data
                    self.encoded_feats[self.label_encoder] = {feature: encoded_data}
                else:
                    if data_types[feature] == 'ordinal':
                        ordinal_mapping = {}
                        print('\nIt is best practice to have the lowest ordered as 0 and go upwards.\n')
                        i = 0
                        while i < len(self.categories):
                            order = input(f'What is the rank of {self.categories[i]}: ')
                            order = self.helper_functions.convert_to_num(order)
                            ordinal_mapping[order] = self.categories[i] #storing the order with the categorical variable in dict
                            if i == len(self.categories) - 1:
                                unique_order_values = set(list(ordinal_mapping.keys()))
                                if len(unique_order_values) < len(self.categories):
                                    i = -1
                            i+=1
                        
                        encoded_data = self.ordinal_mapping(feature, ordinal_mapping)
                        data[feature] = encoded_data
                        self.encoded_feats[self.ordinal_encoder] = {feature: encoded_data}
                        
                    else:
                        
                        data_to_encode = self.helper_functions.reshape_array(data[feature])
                       
                        encoded_data = self.one_hot_encoder.fit_transform(data[[feature]])
                        encoded_df = pd.DataFrame(encoded_data.toarray(), columns=self.one_hot_encoder.get_feature_names_out())
                        self.encoded_feats[self.one_hot_encoder] = {feature: encoded_df}
                        data = pd.concat([data.drop(feature, axis=1), encoded_df], axis=1)

            else:
                title = 'Dataframe Not Encoded:'


        print('\n'+title+'\n')
        print(data.head())

        if self.encoded_feats:
            print('\nOriginal Dataframe:\n')
            print(self.data.head())

        # if not self.encoded_feats:
        #     return self.data
        # else:
        return data, self.encoded_feats                         
                                                      
    def find_null_values(self):
        # print('\033[1m' + 'Finding Null Values\n' + '\033[0m')
        self.helper_functions.format_titles('finding null values')

        total_null_values = self.data.isnull().sum().sum()

        if total_null_values > 0:
            print(f'Total number of null values: {total_null_values}')
            df_null_mask = self.data.isnull()
            self.figure_manager.make_heatmap(df_null_mask, type='dataset_missing', x_label='Columns', y_label='Rows')
            for feature in self.data.columns:
                null_values = self.data[feature].isnull().sum()
                percent_null = (null_values/(int(self.data.shape[0]))*100)
                if null_values > 0:
                    print(f'Number of null values for {feature}: {null_values}')
                    null_mask = self.data[feature]
                    self.figure_manager.make_heatmap(null_mask, type='feat_missing', x_label='Columns', y_label='Rows', feature=feature)
                else:
                    pass
        else:
            print('No Missing values in the dataset')
              
    def find_duplicate_values(self):
        # print('\033[1m' + '\nFinding Duplicated Values\n' + '\033[0m')
        self.helper_functions.format_titles('finding duplicate values')
        duplicated_rows = self.data[self.data.duplicated(keep=False)]
        print(f'Total Duplicated Rows: {len(duplicated_rows)}')
        #duplicated_df = pd.DataFrame(duplicated_rows, columns = self.data.columns)

        
        if len(duplicated_rows) > 0:
            print('\nPotential Duplicates:\n')
            print(duplicated_rows)
            remove = input('\nWould you like to remove the duplicates? Y/N ')
            #remove = self.helper_functions.while_string_function(remove)
            if self.helper_functions.while_string_function(remove):
                print('\nRemoving duplicates')
                new_data = self.data.drop_duplicates()
                if new_data is not self.data:
                    print('Duplicates removed')
                    self.data = new_data
                else:
                    print('Duplicates not removed')
            else:
                pass
           
class Analysis:
    def __init__(self, data, project, target=None):
        # initialising class variables 
        self.data = data.copy()
        self.project = project
        self.target = target

        # defining other classes used here 
        self.helper_functions = Helper_Functions(project)
        self.figure_manager = Figure_Manager(project)
        self.stats = Statistics(data, project)
        self.preprocess = Pre_Processing(data, project) 

        self.data_types = {} # used to automate the encoding 
        self.prev_norm = None # dictionary of features indicating liklihood of normality based on variety of tests

        self.encoded_dataframe = None
        self.heatmap = True

    def begin_analysis(self):
        self.data_types = self.stats.determine_data_type(data)
        self.preprocess.find_null_values()
        self.preprocess.find_duplicate_values()
        self.distribution_analysis()
                    
    def distribution_analysis(self, data=None, transformations = None, dimension_red = None, polynomial_features = False):
        #check for normality before deciding which outlier tests to use - need to check linearity and normality.
        if data:
            data = data
        else:
            data = self.data

        title = 'beginning distribution analysis'
        if transformations and not polynomial_features:
            title += f'after {transformations}'
        elif not transformations and polynomial_features:
            title += 'after creation of polynomial features'
        elif transformations and polynomial_features:
            title += f'after {transformations}, creation of polynomial features'
        
        if dimension_red:
            dimension_red_name = dimension_red.__class__.__name__
            title += f'and {dimension_red_name}'

        self.helper_functions.format_titles(title)
        print('creating violin plots'.title())
        self.figure_manager.make_violin_plots(data, transformations=transformations)
        print('creating box plots'.title())
        self.figure_manager.make_box_plot(self.data)

        if self.heatmap:
            print('creating heatmap'.title())
            self.encoded_dataframe = self.figure_manager.make_heatmap(data, type='dataset_correlation', data_types=self.data_types)
            self.heatmap = False

        cat_feats = []
        num_data = self.data.select_dtypes(include=(int, float))
        skew_kurtosis = self.stats.skew_kurtosis(num_data)
        norm_result = self.stats.normality_checks(num_data, skew_kurtosis)
        print('creating histograms'.title())
        print('creating quantile-quantile plots'.title())
        print('creating bar charts'.title()) 
        print('creating box plots'.title())
        for feature in data.columns:
            if data[feature].dtype != 'object':
                self.helper_functions.format_titles(f'full results for {feature}')
                num_data = data[feature]
                self.figure_manager.make_histogram(num_data, feature)
                self.figure_manager.make_q_q(num_data, feature)
                feat_norm_result = norm_result[feature]
                skew = skew_kurtosis[feature]['skew']
                kurtosis = skew_kurtosis[feature]['kurtosis']
                print(f'{feat_norm_result}')

            else:
                self.figure_manager.make_bar(self.data[feature], feature)
                num_data = self.data.select_dtypes(include=(float, int))
                cat_feats.append(feature)
                cat_vars = set(self.data[feature])
                for cat_var in cat_vars:
                    data = self.data[self.data[feature] == cat_var]
                    num_data = data.select_dtypes(include=(float, int))
                    self.figure_manager.make_box_plot(data=data, feature=feature, cat_var=cat_var)

        norm_assumptions = self.stats.determine_normality(num_data, norm_result)
        norm_assumptions = self.helper_functions.check_assump_correct(norm_assumptions)
        self.outlier_detection(norm_assumptions, cat_feats = cat_feats)

    def multi_var_analysis(self, corr_matrix):
        pass
        
    def outlier_detection(self, norm_assumptions, cat_feats = [], transformations = None):
        # will be fed results of dist_analysis to then find outliers using 2/3 techniques for normal or non-normal data
        # will split data into categories for outlier detection and send each type of outlier removed datatype to stats test to see if 
        # normality has improved.
        title = 'detecting outliers'
        if transformations:
            title += f'after {transformations}'
        self.helper_functions.format_titles(title)

        def z_test(data):
            return np.where(np.abs(stats.zscore(data)) > 3)
        
        def grubbs_test(data):
            pass

        def norm_assump_tests(data, feature_distribution):
            pass

        if len(cat_feats) > 0:
            cat_outlier_test = input('Would you like to group the data by categorical features and test for outliers by category? Y/N ')
            cat_outlier_test = self.helper_functions.while_string_function(cat_outlier_test)
            if cat_outlier_test and len(cat_feats) > 1:
                cat_feats = input(f'Select which features out of {cat_feats} you would like to group by? ')
            elif cat_outlier_test and len(cat_feats) == 1:
                pass
            else:
                print('Using ungrouped dataset')

            if cat_feats:
                print(f'Grouping data by {cat_feats}')

            for feature, norm_assumption in norm_assumptions.items():
                if 'likely' in norm_assumption.lower():
                    pass
                elif 'unlikely' in norm_assumption.lower():
                    pass
                else:
                    pass 


            pass

        def feature_scale_analysis(self):
            pass
        #find out if scale will impact further analysis - this can then be linked to a method in the manipulation class or perhaps would be better in the helper functions class?


    def collinearity_detection(self):
        corr_data = self.helper_functions.select_numeric_cols(self.data)
        corr_matrix = corr_data.corr()
        self.helper_functions.make_heatmap(corr_matrix, plot_name='correlation heatmap')
        #will include the correlation matrix and VIF those with high collinearity can be sent to dimension reduction.
        pass

class Manipulation(Analysis):
    def __init__(self, data, project, target):
        super().__init__(data, project, target)

    def transforms(self, skew):
        i = 1           #starting from 1 because I will test each feature for normality before transforming (if necessary)
        
        # if skew < 0: #spreads out low numbers to higher
        #     while i < 6:
        #         data_to_power = data**i
        #         i += 1    
        
        j = 1
        # elif skew > 0: #spreads out higher to lower 
        #     while i < 6:
        #         data_to_power = data**(-i)
        #         i += 1
        #     while j < 6:
        #         data_to_root = np.power(data, 1/j)
        pass

            
            

    def dimensionality_reduction(self):
        pass

file_path = '/Users/harrisonclarke/Desktop/Coding/new_eda_test/IRIS.csv'
data = pd.read_csv(file_path)
# preprocessing = Pre_Processing(data, project = 'test', target=None)
#encoded_data = preprocessing.encoder()

analyse_data = Analysis(data, 'test', target = None)
analyse_data.begin_analysis()
